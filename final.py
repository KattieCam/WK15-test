"""
Kattie Cameron-Ervin ITEC 1150-60. This is my Random Recipe Cookbook, final coding that's going to generate a
 few random taco recipes, as well a random photo of a taco.
"""

import requests
from PIL import Image, ImageDraw, ImageFont
import docx

im = Image.open('tacos-unsplash.jpg')  # this open my image
width = im.width  # requesting the height and width of my original photo, and storing data
height = im.height

print(width, height)  # preserving the aspect ratio

half_height = int(height / 2)  # i am reducing the size of my image
half_width = int(width / 2)  # i am saving the changes to my resized images

half_size = im.resize((half_width, half_height))
half_size.show()  # this would show my half size image
half_size.save('half_size_taco.jpg')

font = ImageFont.truetype('DejaVuSans.ttf', size=250)  # font = ImageFont.truetype('DejaVuSans.ttf', 50)
draw = ImageDraw.Draw(im)  # drawing tool
draw.text((50, 50), "Random Taco Cookbook", (225, 225, 0), font=font)  # specifiying my font color

im.show()

im.save('half_size_taco_text.jpg')

taco_data = 'https://taco-1150.herokuapp.com/random/?full_taco=true'  # requesting data from url
response = requests.get(taco_data)

if not response:
    print('page you requested cannot be found')  # checking if data is available
    quit()

response_json = response.json()  # this will bring back a dictionary of variables, of keys and ingredients

print('printing the name of the seasoning:')  # returning name of seasoning
print(response_json['seasoning']['name'])
print('printing the recipe of the seasoning:')
print(response_json['seasoning']['recipe'])  # return recipe to seasoning

# use this if you want to view all of the data / to see what the keys are
"""
for first_key in response_json.keys(): 
    item = response_json[first_key]
    for second_key in item.keys()
    print(f'--------------------------------')
    print(f'[{first_key}][{second_key}]')
    print(f'--------------------------------')
    print(item[second_key])
    print()
"""


# this section is where i was generate

def create_document(half_size, taco_data):
    document = docx.Document()
    # naming and adding heading
    document.add_heading("Random Taco Cookbook", 0)
    # adding a picture
    document.add_picture(half_size, width=(width / 2))
    # adding and naming heading for credits
    document.add_heading('Credits')
    # adding lines under credits
    document.add_paragraph('Taco image: ' + taco_data, )
    document.add_paragraph('Tacos from: ' + taco_data, )
    document.add_paragraph('Created by: Kattie')

    for recipe in taco_data:  # loops each recipe in taco recipes
        document.add_page_break()
        # adds a page break
        names = ', '.join([name['name'] for name in recipe])
        # creates a headline for each recipe
        document.add_heading(names, 0)

    for rec in recipe:
        document.add_heading(recipe[rec]['name'], level=1)  # creates a heading for recipe list
        document.add_paragraph(recipe[rec]['recipe'])  # adds list of recipe
        # after looping, save taco data in a word document
        document.save('recipes.docx')
